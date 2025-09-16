# АБГВА
# А: Сгенерировать синтетические данные для 10 строк (формат: имя, количество, цена)
# Б: Добавить условное форматирование: выделить ячейки в колонке "Количество" цветом, если значение больше 100
# Г: Создать документ с двумя разделами: первый - исходные данные (таблица), второй - итоги (итоговая сумма и самая дорогая позиция)
# В: Создать PDF, вставив в него изображение (например, логотип) и итоговую текстовую сводку
# A: Сохранить все файлы (XLSX, DOCX, PDF) в текущую директорию

import random
from openpyxl import Workbook
from openpyxl.styles import PatternFill
from openpyxl.formatting.rule import CellIsRule
from docx import Document
from fpdf import FPDF
from fpdf.enums import XPos, YPos

names = ["Карандаш", "Ручка", "Блокнот", "Тетрадь", "Стиральная резинка", "Линейка", "Фломастеры", "Маркер", "Папка", "Клей"]

data = []
for name in names:
    quantity = random.randint(1, 200)
    price = random.randint(50, 500)
    data.append((name, quantity, price))

##############################################################################################

wb = Workbook()
ws = wb.active
ws.title = "Данные"
ws.append(["Товар", "Количество", "Цена"])
for row in data:
    ws.append(row)

yellow_fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')
ws.conditional_formatting.add(
    'B2:B11',
    CellIsRule(operator='greaterThan', formula=['100'], fill=yellow_fill)
)
excel_filename = "Lab #2.xlsx"
wb.save(excel_filename)

##############################################################################################

doc = Document()
doc.add_heading("Отчёт", 0)
doc.add_heading("1. Исходные данные", level=1)
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Товар'
hdr_cells[1].text = 'Количество'
hdr_cells[2].text = 'Цена'

for name, quantity, price in data:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = str(quantity)
    row_cells[2].text = str(price)

total_sum = sum(quantity * price for _, quantity, price in data)
most_expensive = max(data, key=lambda x: x[2])

doc.add_heading("2. Итоги", level=1)
doc.add_paragraph(f"Итоговая сумма: {round(total_sum, 2)}")
doc.add_paragraph(f"Самая дорогая позиция: {most_expensive[0]} ({most_expensive[2]} за единицу)")

word_filename = "Lab #2.docx"
doc.save(word_filename)

##############################################################################################

pdf = FPDF()
pdf.add_page()

font_path = "./Font/ARIAL.TTF"
pdf.add_font("Arial", "", font_path)
pdf.set_font("Arial", size=14)

logo = "logo.png"
pdf.image(logo, x=10, y=8, w=50)
pdf.ln(60)

pdf.cell(200, 10, text="Итоговая текстовая сводка", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
pdf.cell(200, 10, text=f"Итоговая сумма: {round(total_sum, 2)}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
pdf.cell(200, 10, text=f"Самая дорогая позиция: {most_expensive[0]} ({most_expensive[2]} за единицу)", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

pdf_filename = "Lab #2.pdf"
pdf.output(pdf_filename)